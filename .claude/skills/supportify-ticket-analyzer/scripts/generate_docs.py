#!/usr/bin/env python3
"""
Documentation Generator for Supportify
Creates documentation in multiple formats (Markdown, DOCX, PDF) from ticket analysis
"""

import json
import sys
from pathlib import Path
from datetime import datetime

def load_analysis(analysis_file):
    """Load the JSON analysis results"""
    with open(analysis_file, 'r') as f:
        return json.load(f)

def generate_markdown(analysis, output_file):
    """Generate Markdown documentation"""
    summary = analysis['summary']
    apple_tickets = analysis['apple_addressable_tickets']
    
    md_content = f"""# Apple Support Ticket Analysis Report

**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
**Analysis Approach:** Phase 1 - Inclusive categorization prioritizing Apple documentation

## Executive Summary

- **Total Tickets Analyzed:** {summary['total_tickets']}
- **Apple-Addressable Issues:** {summary['apple_addressable']} ({summary['apple_addressable_pct']}%)
- **Vendor-Specific Issues (Limited Apple Docs):** {summary.get('vendor_specific', summary.get('enterprise_it', 0))}
- **Uncategorized (Needs Review):** {summary['uncategorized']}

---

## Apple-Addressable Issues by Category

The following issues can be resolved using Apple's official documentation, support articles, and guidance:

"""
    
    # Add category breakdown
    for category, count in summary['top_apple_categories'].items():
        pct = round(count / summary['apple_addressable'] * 100, 1) if summary['apple_addressable'] > 0 else 0
        md_content += f"\n### {category.replace('_', ' ').title()}\n"
        md_content += f"**Frequency:** {count} tickets ({pct}% of Apple-addressable issues)\n\n"
        
        # Add sample tickets for this category
        category_tickets = [t for t in apple_tickets if t['category'] == category][:3]
        if category_tickets:
            md_content += "**Sample Tickets:**\n\n"
            for ticket in category_tickets:
                md_content += f"- **{ticket['ticket_id']}**: {ticket['short_description']}\n"
                if ticket['description']:
                    md_content += f"  - {ticket['description'][:150]}...\n"
            md_content += "\n"
        
        # Add Apple documentation recommendations
        md_content += "**Recommended Apple Resources:**\n\n"
        md_content += get_apple_resources(category)
        md_content += "\n---\n"
    
    # Add appendix with all Apple-addressable tickets
    md_content += "\n## Appendix: All Apple-Addressable Tickets\n\n"
    md_content += "| Ticket ID | Category | Description |\n"
    md_content += "|-----------|----------|-------------|\n"
    
    for ticket in apple_tickets:
        desc = ticket['description'][:100].replace('|', '-') if ticket['description'] else ''
        md_content += f"| {ticket['ticket_id']} | {ticket['category']} | {desc}... |\n"
    
    # Write to file
    with open(output_file, 'w') as f:
        f.write(md_content)
    
    return output_file

def get_apple_resources(category):
    """Return recommended Apple documentation for each category"""
    resources = {
        'macos_update': """
- [Update macOS on Mac - Apple Support](https://support.apple.com/guide/mac-help/update-macos-mchlpx1065)
- [If your Mac doesn't recognize software updates - Apple Support](https://support.apple.com/HT211683)
- [macOS Sequoia is compatible with these computers - Apple Support](https://support.apple.com/HT213264)
- [Prepare your Mac for a macOS update](https://support.apple.com/guide/mac-help/prepare-your-mac-macos-update-mh15915)
""",
        'login_issue': """
- [If you can't log in to your Mac with your password - Apple Support](https://support.apple.com/HT202860)
- [Use FileVault to encrypt the startup disk - Apple Support](https://support.apple.com/guide/mac-help/protect-data-filevault-mh11785)
- [Manage keychain passwords - Apple Support](https://support.apple.com/guide/mac-help/keychain-access-mchlf375f392)
- [Reset your Mac login password - Apple Support](https://support.apple.com/HT212190)
""",
        'hardware_issue': """
- [Apple Diagnostics - Apple Support](https://support.apple.com/HT202731)
- [Mac notebook battery information - Apple Support](https://support.apple.com/mac-battery)
- [Find the serial number or IMEI on your Mac - Apple Support](https://support.apple.com/HT201581)
- [Service and repair for Mac - Apple Support](https://support.apple.com/mac/repair)
""",
        'wifi_network': """
- [Connect to Wi-Fi on your Mac - Apple Support](https://support.apple.com/guide/mac-help/connect-to-wi-fi-mchlp2550)
- [If your Mac can't connect to the internet over Wi-Fi - Apple Support](https://support.apple.com/HT202663)
- [Use your Mac as a Wi-Fi hotspot - Apple Support](https://support.apple.com/guide/mac-help/share-internet-connection-mac-network-mchlp1540)
""",
        'vpn_issue': """
- [Set up a VPN connection on Mac - Apple Support](https://support.apple.com/guide/mac-help/set-up-vpn-connection-mchlp2963)
- [Change VPN settings on Mac - Apple Support](https://support.apple.com/guide/mac-help/change-vpn-settings-mchlp2322)
""",
        'performance': """
- [If your Mac runs slowly - Apple Support](https://support.apple.com/HT204129)
- [Check storage on your Mac - Apple Support](https://support.apple.com/HT206996)
- [Free up storage space on your Mac - Apple Support](https://support.apple.com/HT206996)
- [If an app freezes or quits unexpectedly - Apple Support](https://support.apple.com/HT203004)
""",
        'native_apps': """
- [Mail User Guide for Mac - Apple Support](https://support.apple.com/guide/mail)
- [Safari User Guide for Mac - Apple Support](https://support.apple.com/guide/safari)
- [Messages User Guide for Mac - Apple Support](https://support.apple.com/guide/messages)
""",
        'bluetooth': """
- [Connect a Bluetooth device with your Mac - Apple Support](https://support.apple.com/HT204350)
- [Use AirDrop on your Mac - Apple Support](https://support.apple.com/HT203106)
- [If your Mac doesn't recognize Bluetooth accessories - Apple Support](https://support.apple.com/HT201154)
""",
        'printer': """
- [Add a printer on Mac - Apple Support](https://support.apple.com/guide/mac-help/add-printer-mh14004)
- [Use AirPrint to print from your Mac - Apple Support](https://support.apple.com/HT201311)
- [If you can't print from your Mac - Apple Support](https://support.apple.com/HT5128)
""",
        'permissions': """
- [Change Privacy & Security settings on Mac - Apple Support](https://support.apple.com/guide/mac-help/change-privacy-security-settings-mh40596)
- [Control app access to files and folders on Mac - Apple Support](https://support.apple.com/guide/mac-help/control-access-files-folders-mchld5a35146)
""",
        'mdm_management': """
- [Apple Platform Deployment - Apple Support](https://support.apple.com/guide/deployment/)
- [Intro to Apple Business Manager - Apple Support](https://support.apple.com/guide/apple-business-manager/)
- [Device enrollment overview - Apple Support](https://support.apple.com/guide/deployment/intro-to-mdm-dep00002)
- [Configuration Profile Reference - Apple Developer](https://developer.apple.com/documentation/devicemanagement/configuring_multiple_devices)
- [Use Mobile Device Management - Apple Support](https://support.apple.com/guide/deployment/use-mobile-device-management-dep0dd0c7e7)
""",
        'app_deployment': """
- [Deploy apps on Mac - Apple Support](https://support.apple.com/guide/deployment/deploy-apps-depc0daa4f58)
- [Manage apps with MDM - Apple Support](https://support.apple.com/guide/deployment/manage-apps-with-mdm-dep93a3baa3)
- [Create a package installer for Mac - Apple Developer](https://developer.apple.com/documentation/xcode/distributing-software-using-installer-packages)
- [About app installation on Mac - Apple Support](https://support.apple.com/guide/mac-help/install-and-uninstall-apps-mh35835)
""",
        'enterprise_auth': """
- [Integrate macOS with Microsoft Active Directory - Apple Support](https://support.apple.com/guide/deployment/integrate-with-active-directory-dep9e693b01)
- [Use Kerberos for single sign-on - Apple Support](https://support.apple.com/guide/deployment/use-kerberos-for-single-sign-on-dep35bc5e9a)
- [Network account server settings on Mac - Apple Support](https://support.apple.com/guide/mac-help/network-account-server-settings-mh15410)
- [Certificate-based authentication - Apple Support](https://support.apple.com/guide/deployment/certificate-based-authentication-dep10a9bb77)
""",
        'enterprise_network': """
- [Configure network settings on Mac - Apple Support](https://support.apple.com/guide/deployment/configure-network-settings-depbc44d9bb)
- [Set up 802.1X authentication - Apple Support](https://support.apple.com/guide/deployment/set-up-8021x-authentication-dep5712a6e6)
- [Configure proxy settings on Mac - Apple Support](https://support.apple.com/guide/mac-help/change-proxy-settings-mchlp2591)
- [Install certificates on Mac - Apple Support](https://support.apple.com/guide/deployment/install-certificates-depc0cb3c77a)
- [VPN configurations for iPhone and iPad - Apple Support](https://support.apple.com/guide/deployment/vpn-configurations-depf02ad364)
""",
        'file_sharing': """
- [Set up file sharing on Mac - Apple Support](https://support.apple.com/guide/mac-help/set-up-file-sharing-mh17131)
- [Connect to a server or shared computer - Apple Support](https://support.apple.com/guide/mac-help/connect-mac-shared-computers-servers-mchlp1140)
- [Change permissions for files, folders, or disks - Apple Support](https://support.apple.com/guide/mac-help/change-permissions-files-folders-disks-mchlp1203)
- [Share files between a Mac and Windows PC - Apple Support](https://support.apple.com/guide/mac-help/share-files-between-mac-windows-mchlp1657)
""",
        'software_update_mgmt': """
- [Manage software updates on Mac - Apple Support](https://support.apple.com/guide/deployment/manage-software-updates-depc4c80847a)
- [About software update downloads - Apple Support](https://support.apple.com/HT211683)
- [Use managed software updates - Apple Support](https://support.apple.com/guide/deployment/use-managed-software-updates-depc24465c0c)
- [macOS update catalogs - Apple Support](https://support.apple.com/guide/deployment/macos-update-catalogs-depc331908f1)
""",
        'device_lifecycle': """
- [Set up your new Mac - Apple Support](https://support.apple.com/guide/mac-help/set-up-your-new-mac-mchl8cf3c48c)
- [Move your data to a new Mac - Apple Support](https://support.apple.com/HT204350)
- [What to do before selling, giving away, or trading in your Mac - Apple Support](https://support.apple.com/HT201065)
- [Erase your Mac and reset it to factory settings - Apple Support](https://support.apple.com/HT212749)
- [Remove Activation Lock on Mac - Apple Support](https://support.apple.com/HT208987)
- [Use Apple Configurator to provision Mac computers - Apple Support](https://support.apple.com/guide/apple-configurator-mac/)
"""
    }

    return resources.get(category, "- Search Apple Support articles at https://support.apple.com\n")

def generate_docx(analysis, output_file):
    """Generate DOCX documentation"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Apple Support Ticket Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = doc.add_paragraph()
        date_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Space
        
        # Executive Summary
        doc.add_heading('Executive Summary', 1)
        summary = analysis['summary']
        
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = 'Light Grid Accent 1'
        
        summary_table.rows[0].cells[0].text = 'Total Tickets Analyzed'
        summary_table.rows[0].cells[1].text = str(summary['total_tickets'])
        
        summary_table.rows[1].cells[0].text = 'Apple-Addressable Issues'
        summary_table.rows[1].cells[1].text = f"{summary['apple_addressable']} ({summary['apple_addressable_pct']}%)"

        summary_table.rows[2].cells[0].text = 'Vendor-Specific Issues (Limited Apple Docs)'
        summary_table.rows[2].cells[1].text = str(summary.get('vendor_specific', summary.get('enterprise_it', 0)))

        summary_table.rows[3].cells[0].text = 'Uncategorized (Needs Review)'
        summary_table.rows[3].cells[1].text = str(summary['uncategorized'])
        
        doc.add_page_break()
        
        # Categories
        doc.add_heading('Apple-Addressable Issues by Category', 1)
        
        apple_tickets = analysis['apple_addressable_tickets']
        
        for category, count in summary['top_apple_categories'].items():
            pct = round(count / summary['apple_addressable'] * 100, 1) if summary['apple_addressable'] > 0 else 0
            
            doc.add_heading(category.replace('_', ' ').title(), 2)
            doc.add_paragraph(f"Frequency: {count} tickets ({pct}% of Apple-addressable issues)")
            
            # Sample tickets
            category_tickets = [t for t in apple_tickets if t['category'] == category][:3]
            if category_tickets:
                doc.add_paragraph('Sample Tickets:', style='Intense Quote')
                for ticket in category_tickets:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f"{ticket['ticket_id']}: ").bold = True
                    p.add_run(f"{ticket['short_description']}")
            
            # Resources
            doc.add_paragraph('Recommended Apple Resources:', style='Intense Quote')
            resources = get_apple_resources(category)
            for line in resources.strip().split('\n'):
                if line.strip().startswith('-'):
                    doc.add_paragraph(line.strip()[2:], style='List Bullet')
            
            doc.add_paragraph()  # Space
        
        # Save
        doc.save(output_file)
        return output_file
        
    except ImportError:
        print("Warning: python-docx not installed. Install with: pip install python-docx")
        return None

def generate_pdf(markdown_file, output_file):
    """Convert Markdown to PDF using pandoc or weasyprint"""
    import subprocess
    
    try:
        # Try using pandoc first
        result = subprocess.run(
            ['pandoc', markdown_file, '-o', output_file, '--pdf-engine=xelatex'],
            capture_output=True,
            text=True
        )
        if result.returncode == 0:
            return output_file
    except FileNotFoundError:
        pass
    
    # Try weasyprint
    try:
        from markdown import markdown
        from weasyprint import HTML
        
        with open(markdown_file, 'r') as f:
            md_content = f.read()
        
        html_content = f"""
        <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                h1 {{ color: #1d1d1f; border-bottom: 2px solid #0066cc; }}
                h2 {{ color: #0066cc; margin-top: 30px; }}
                table {{ border-collapse: collapse; width: 100%; }}
                th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                th {{ background-color: #0066cc; color: white; }}
            </style>
        </head>
        <body>
            {markdown(md_content, extensions=['tables', 'fenced_code'])}
        </body>
        </html>
        """
        
        HTML(string=html_content).write_pdf(output_file)
        return output_file
        
    except ImportError:
        print("Warning: Neither pandoc nor weasyprint available for PDF generation")
        print("Install with: pip install weasyprint markdown")
        return None

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python generate_docs.py <analysis_json> [output_dir]")
        sys.exit(1)
    
    analysis_file = sys.argv[1]
    output_dir = Path(sys.argv[2]) if len(sys.argv) > 2 else Path('.')
    output_dir.mkdir(exist_ok=True)
    
    # Load analysis
    analysis = load_analysis(analysis_file)
    
    # Generate documentation
    print("Generating documentation...")
    
    # Markdown
    md_file = output_dir / 'support_analysis.md'
    generate_markdown(analysis, md_file)
    print(f"✓ Markdown: {md_file}")
    
    # DOCX
    docx_file = output_dir / 'support_analysis.docx'
    result = generate_docx(analysis, docx_file)
    if result:
        print(f"✓ DOCX: {docx_file}")
    
    # PDF
    pdf_file = output_dir / 'support_analysis.pdf'
    result = generate_pdf(md_file, pdf_file)
    if result:
        print(f"✓ PDF: {pdf_file}")
    
    print("\nDocumentation generation complete!")
