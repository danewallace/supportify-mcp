#!/usr/bin/env python3
"""
Multi-Format KB Article Exporter
Exports KB articles to PDF, DOCX, HTML, and ServiceNow JSON formats
"""

import json
import sys
import re
from pathlib import Path
from datetime import datetime

# Import dependencies
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError):
    # OSError can occur on macOS if system libraries aren't installed
    WEASYPRINT_AVAILABLE = False

def parse_frontmatter(content):
    """Extract YAML frontmatter from markdown"""
    frontmatter = {}
    body = content

    if content.startswith('---'):
        parts = content.split('---', 2)
        if len(parts) >= 3:
            yaml_content = parts[1].strip()
            body = parts[2].strip()

            for line in yaml_content.split('\n'):
                if ':' in line:
                    key, value = line.split(':', 1)
                    key = key.strip()
                    value = value.strip()

                    if value.startswith('[') and value.endswith(']'):
                        value = [v.strip() for v in value[1:-1].split(',')]

                    frontmatter[key] = value

    return frontmatter, body

def markdown_to_html(markdown_text, full_page=False):
    """Convert markdown to HTML"""
    if MARKDOWN_AVAILABLE:
        # Pre-process: Convert checkmark lines to list items
        # Lines starting with ✓ should become list items
        processed_text = []
        lines = markdown_text.split('\n')
        in_checkmark_list = False

        for i, line in enumerate(lines):
            if line.strip().startswith('✓ '):
                if not in_checkmark_list:
                    # Start of checkmark list
                    in_checkmark_list = True
                # Convert ✓ to - for markdown list
                processed_text.append('- ' + line.strip()[2:])
            else:
                if in_checkmark_list and line.strip() and not line.strip().startswith('✓'):
                    # End of checkmark list
                    in_checkmark_list = False
                processed_text.append(line)

        markdown_text = '\n'.join(processed_text)

        html_body = markdown.markdown(
            markdown_text,
            extensions=['extra', 'codehilite', 'tables', 'toc']
        )

        # Post-process: Style checkmark list items
        html_body = html_body.replace('<li>✓', '<li class="checkmark">✓')
        html_body = re.sub(r'<li>([^<]*)</li>', lambda m: f'<li>{m.group(1)}</li>' if not m.group(1).strip().startswith('✓') else f'<li class="checkmark">{m.group(1)}</li>', html_body)
    else:
        # Basic fallback
        html_body = markdown_text
        html_body = re.sub(r'^# (.*?)$', r'<h1>\1</h1>', html_body, flags=re.MULTILINE)
        html_body = re.sub(r'^## (.*?)$', r'<h2>\1</h2>', html_body, flags=re.MULTILINE)
        html_body = re.sub(r'^### (.*?)$', r'<h3>\1</h3>', html_body, flags=re.MULTILINE)
        html_body = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', html_body)
        html_body = re.sub(r'\*(.*?)\*', r'<em>\1</em>', html_body)
        html_body = re.sub(r'```(.*?)```', r'<pre><code>\1</code></pre>', html_body, flags=re.DOTALL)

    if full_page:
        return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>KB Article</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
            line-height: 1.6;
            max-width: 900px;
            margin: 40px auto;
            padding: 20px;
            color: #333;
        }}
        h1 {{
            color: #1d1d1f;
            border-bottom: 3px solid #0066cc;
            padding-bottom: 10px;
            margin-top: 30px;
        }}
        h2 {{
            color: #0066cc;
            margin-top: 30px;
            border-bottom: 1px solid #e0e0e0;
            padding-bottom: 5px;
        }}
        h3 {{
            color: #333;
            margin-top: 20px;
        }}
        code {{
            background-color: #f5f5f7;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: 'SF Mono', Monaco, 'Courier New', monospace;
            font-size: 0.9em;
        }}
        pre {{
            background-color: #f5f5f7;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
        }}
        pre code {{
            background-color: transparent;
            padding: 0;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }}
        th {{
            background-color: #0066cc;
            color: white;
        }}
        tr:nth-child(even) {{
            background-color: #f9f9f9;
        }}
        blockquote {{
            border-left: 4px solid #0066cc;
            padding-left: 20px;
            margin-left: 0;
            color: #666;
        }}
        a {{
            color: #0066cc;
            text-decoration: none;
        }}
        a:hover {{
            text-decoration: underline;
        }}
        hr {{
            border: none;
            border-top: 2px solid #e0e0e0;
            margin: 30px 0;
        }}
        .metadata {{
            background-color: #f0f0f0;
            padding: 15px;
            border-radius: 5px;
            margin-bottom: 20px;
            font-size: 0.9em;
        }}
        .checkmark {{
            color: #00AA00;
            font-weight: 500;
        }}
        li.checkmark {{
            margin-bottom: 8px;
        }}
    </style>
</head>
<body>
{html_body}
</body>
</html>"""
    return html_body

def export_html(markdown_file, output_file, frontmatter=None):
    """Export to HTML format"""
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()

    fm, body = parse_frontmatter(content)
    html_content = markdown_to_html(body, full_page=True)

    # Add metadata section if frontmatter exists
    if fm:
        metadata_html = '<div class="metadata">'
        if 'title' in fm:
            metadata_html += f'<strong>Title:</strong> {fm["title"]}<br>'
        if 'category' in fm:
            metadata_html += f'<strong>Category:</strong> {fm["category"]}<br>'
        if 'frequency' in fm:
            metadata_html += f'<strong>Based on:</strong> {fm["frequency"]}<br>'
        if 'source' in fm:
            metadata_html += f'<strong>Source:</strong> {fm["source"]}<br>'
        metadata_html += '</div>'
        html_content = html_content.replace('<body>', f'<body>\n{metadata_html}')

    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(html_content)

    return output_file

def export_docx(markdown_file, output_file):
    """Export to DOCX format"""
    if not DOCX_AVAILABLE:
        print("Warning: python-docx not available. Install with: pip install python-docx")
        return None

    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()

    fm, body = parse_frontmatter(content)

    doc = Document()

    # Set document style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Helvetica'
    font.size = Pt(11)

    # Add title
    if 'title' in fm:
        title = doc.add_heading(fm['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(0, 102, 204)

    # Add metadata
    metadata_table = doc.add_table(rows=0, cols=2)
    metadata_table.style = 'Light Grid Accent 1'

    if 'category' in fm:
        row = metadata_table.add_row()
        row.cells[0].text = 'Category'
        row.cells[1].text = str(fm['category'])

    if 'frequency' in fm:
        row = metadata_table.add_row()
        row.cells[0].text = 'Based on'
        row.cells[1].text = str(fm['frequency'])

    if 'source' in fm:
        row = metadata_table.add_row()
        row.cells[0].text = 'Source'
        row.cells[1].text = str(fm['source'])

    if 'created' in fm:
        row = metadata_table.add_row()
        row.cells[0].text = 'Created'
        row.cells[1].text = str(fm['created'])

    doc.add_paragraph()  # Space after metadata

    # Helper function to add proper bullet list items
    def add_bullet_paragraph(text, level=0):
        """Add a properly formatted bullet list item with actual bullets"""
        para = doc.add_paragraph(style='List Bullet')

        # Add proper numbering format for bullets
        pPr = para._element.get_or_add_pPr()
        numPr = OxmlElement('w:numPr')

        # Bullet list uses numId=1, ilvl for indentation level
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), '1')

        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), str(level))

        numPr.append(ilvl)
        numPr.append(numId)
        pPr.append(numPr)

        return para

    # Helper function to add proper numbered list items
    def add_numbered_paragraph(text, level=0):
        """Add a properly formatted numbered list item"""
        para = doc.add_paragraph(style='List Number')

        # Add proper numbering format
        pPr = para._element.get_or_add_pPr()
        numPr = OxmlElement('w:numPr')

        # Numbered list uses numId=2, ilvl for indentation level
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), '2')

        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), str(level))

        numPr.append(ilvl)
        numPr.append(numId)
        pPr.append(numPr)

        return para

    # Helper function to process inline formatting
    def process_inline_formatting(text, paragraph):
        """Process bold, italic, code, and links in text"""
        # Patterns for different formatting
        link_pattern = r'\[([^\]]+)\]\(([^\)]+)\)'
        bold_pattern = r'\*\*(.*?)\*\*'
        italic_pattern = r'\*([^*]+?)\*'
        code_pattern = r'`([^`]+)`'

        # Combine all patterns (without capturing groups in outer level)
        combined_pattern = f'({link_pattern}|{bold_pattern}|{italic_pattern}|{code_pattern})'

        last_end = 0

        for match in re.finditer(combined_pattern, text):
            # Add any text before this match
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])

            matched_text = match.group(0)

            # Check what type of formatting this is
            if matched_text.startswith('['):
                # Link
                link_match = re.match(link_pattern, matched_text)
                if link_match:
                    link_text = link_match.group(1)
                    link_url = link_match.group(2)

                    # Skip local file paths
                    if link_url.startswith('/Users/') or link_url.startswith('file://'):
                        run = paragraph.add_run(link_text)
                        run.font.color.rgb = RGBColor(0, 102, 204)
                    else:
                        add_hyperlink(paragraph, link_url, link_text)

            elif matched_text.startswith('**'):
                # Bold
                run = paragraph.add_run(matched_text[2:-2])
                run.bold = True

            elif matched_text.startswith('`'):
                # Code
                run = paragraph.add_run(matched_text[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(80, 80, 80)

            elif matched_text.startswith('*'):
                # Italic
                run = paragraph.add_run(matched_text[1:-1])
                run.italic = True

            last_end = match.end()

        # Add any remaining text after the last match
        if last_end < len(text):
            paragraph.add_run(text[last_end:])

    def add_hyperlink(paragraph, url, text):
        """Add a hyperlink-styled run to a paragraph"""
        # python-docx hyperlink creation is complex and error-prone
        # Using styled text (blue + underline) provides clear visual indication
        run = paragraph.add_run(text)
        run.font.color.rgb = RGBColor(0, 102, 204)  # Blue color
        run.underline = True

        # Note: The URL is intentionally not displayed - this is standard for
        # hyperlink text in documents. Users know blue underlined text is clickable.

    # Process markdown content
    lines = body.split('\n')
    in_code_block = False
    code_block_content = []

    for line in lines:
        # Code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                if code_block_content:
                    code_para = doc.add_paragraph('\n'.join(code_block_content))
                    code_para.style = 'Intense Quote'
                    code_para_format = code_para.paragraph_format
                    code_para_format.left_indent = Inches(0.5)
                code_block_content = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            continue

        if in_code_block:
            code_block_content.append(line)
            continue

        # Headers
        if line.startswith('### '):
            heading_text = line[4:].strip()
            heading_para = doc.add_heading('', 3)
            process_inline_formatting(heading_text, heading_para)
        elif line.startswith('## '):
            heading_text = line[3:].strip()
            heading_para = doc.add_heading('', 2)
            process_inline_formatting(heading_text, heading_para)
        elif line.startswith('# '):
            heading_text = line[2:].strip()
            heading_para = doc.add_heading('', 1)
            process_inline_formatting(heading_text, heading_para)
        # Horizontal rule
        elif line.strip() == '---':
            hr_para = doc.add_paragraph()
            # Add bottom border to create a clean horizontal line
            pPr = hr_para._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')  # Border width (1/8 pt)
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'D0D0D0')  # Light gray
            pBdr.append(bottom)
            pPr.append(pBdr)
            # Add small spacing
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:after'), '100')
            spacing.set(qn('w:before'), '100')
            pPr.append(spacing)
        # Lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            list_text = line.strip()[2:]
            para = add_bullet_paragraph(list_text)
            process_inline_formatting(list_text, para)
        elif re.match(r'^\d+\. ', line.strip()):
            list_text = re.sub(r'^\d+\. ', '', line.strip())
            para = add_numbered_paragraph(list_text)
            process_inline_formatting(list_text, para)
        # Checkmarks
        elif line.strip().startswith('✓ '):
            para = doc.add_paragraph()
            run = para.add_run('✓ ')
            run.font.color.rgb = RGBColor(0, 153, 0)
            run.bold = True
            process_inline_formatting(line.strip()[2:], para)
        # Regular paragraphs
        elif line.strip():
            para = doc.add_paragraph()
            process_inline_formatting(line, para)

    # Save document
    doc.save(output_file)
    return output_file

def export_pdf(markdown_file, output_file):
    """Export to PDF format using WeasyPrint"""
    if not WEASYPRINT_AVAILABLE:
        print("Warning: weasyprint not available. Install with: pip install weasyprint")
        return None

    # First create HTML
    html_file = output_file.replace('.pdf', '_temp.html')
    export_html(markdown_file, html_file)

    # Convert HTML to PDF
    try:
        HTML(html_file).write_pdf(output_file)
        # Clean up temp HTML
        Path(html_file).unlink()
        return output_file
    except Exception as e:
        print(f"Error generating PDF: {e}")
        return None

def export_all_formats(markdown_file, output_dir, base_name=None):
    """
    Export a single markdown KB article to all formats
    """
    md_path = Path(markdown_file)
    out_path = Path(output_dir)
    out_path.mkdir(parents=True, exist_ok=True)

    if base_name is None:
        base_name = md_path.stem

    results = {}

    print(f"Exporting: {md_path.name}")
    print(f"Output directory: {out_path}")
    print()

    # 1. Copy markdown (already exists)
    md_output = out_path / f"{base_name}.md"
    import shutil
    shutil.copy(markdown_file, md_output)
    results['markdown'] = str(md_output)
    print(f"✓ Markdown: {md_output.name}")

    # 2. Export HTML
    html_output = out_path / f"{base_name}.html"
    if export_html(markdown_file, html_output):
        results['html'] = str(html_output)
        print(f"✓ HTML: {html_output.name}")

    # 3. Export DOCX
    docx_output = out_path / f"{base_name}.docx"
    if export_docx(markdown_file, docx_output):
        results['docx'] = str(docx_output)
        print(f"✓ DOCX: {docx_output.name}")
    else:
        print(f"✗ DOCX: Not available (install python-docx)")

    # 4. Export PDF
    pdf_output = out_path / f"{base_name}.pdf"
    if export_pdf(markdown_file, pdf_output):
        results['pdf'] = str(pdf_output)
        print(f"✓ PDF: {pdf_output.name}")
    else:
        print(f"✗ PDF: Not available (install weasyprint)")

    # 5. Export ServiceNow JSON
    try:
        # Import the ServiceNow export function
        import sys
        script_dir = Path(__file__).parent
        sys.path.insert(0, str(script_dir))
        from export_servicenow import create_servicenow_kb_json

        json_output = out_path / f"{base_name}_servicenow.json"
        kb_record = create_servicenow_kb_json(markdown_file)
        with open(json_output, 'w', encoding='utf-8') as f:
            json.dump(kb_record, f, indent=2, ensure_ascii=False)
        results['servicenow_json'] = str(json_output)
        print(f"✓ ServiceNow JSON: {json_output.name}")
    except Exception as e:
        print(f"✗ ServiceNow JSON: Error - {e}")

    return results

def export_batch(input_dir, output_dir):
    """Export all KB articles in a directory to all formats"""
    input_path = Path(input_dir)
    kb_files = list(input_path.glob('*_kb*.md')) + list(input_path.glob('*KB*.md'))

    if not kb_files:
        print(f"No KB articles found in {input_dir}")
        return []

    all_results = []

    print(f"Found {len(kb_files)} KB articles to export")
    print("=" * 80)
    print()

    for kb_file in kb_files:
        results = export_all_formats(kb_file, output_dir, kb_file.stem)
        all_results.append({
            'source': str(kb_file),
            'exports': results
        })
        print()

    return all_results

if __name__ == '__main__':
    if len(sys.argv) < 2 or '--help' in sys.argv or '-h' in sys.argv:
        print("""
Multi-Format KB Article Exporter
=================================

Export KB articles to multiple formats: Markdown, HTML, DOCX, PDF, ServiceNow JSON

Usage:
    # Export single article to all formats
    python export_kb_all_formats.py article.md output_dir/

    # Export batch of articles
    python export_kb_all_formats.py --batch kb_templates/ output_dir/

Options:
    --batch     Process all KB articles in directory

Output Formats:
    - Markdown (.md)           Original format
    - HTML (.html)             Web publishing
    - DOCX (.docx)             Microsoft Word (editable)
    - PDF (.pdf)               Distribution/printing
    - ServiceNow JSON (.json)  ServiceNow import

Requirements:
    - python-docx: pip install python-docx
    - markdown: pip install markdown
    - weasyprint: pip install weasyprint

Example:
    # Single article
    python export_kb_all_formats.py wifi_network_kb.md ./exports/

    # Batch export
    python export_kb_all_formats.py --batch ./kb_templates/ ./all_exports/
""")
        sys.exit(0)

    batch_mode = '--batch' in sys.argv

    try:
        if batch_mode:
            args = [arg for arg in sys.argv[1:] if not arg.startswith('--')]
            if len(args) < 2:
                print("Error: Batch mode requires: --batch <input_dir> <output_dir>")
                sys.exit(1)

            input_dir = args[0]
            output_dir = args[1]

            results = export_batch(input_dir, output_dir)

            print("=" * 80)
            print(f"Batch export complete! Exported {len(results)} KB articles")
            print(f"Output directory: {output_dir}")

        else:
            if len(sys.argv) < 3:
                print("Error: Single mode requires: <input.md> <output_dir>")
                sys.exit(1)

            markdown_file = sys.argv[1]
            output_dir = sys.argv[2]

            results = export_all_formats(markdown_file, output_dir)

            print()
            print("=" * 80)
            print("Export complete!")
            print(f"Output directory: {output_dir}")

    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
